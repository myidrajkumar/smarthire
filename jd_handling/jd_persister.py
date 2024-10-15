"""Save the JD"""

import pathlib
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    ListFlowable,
    ListItem,
    Image,
    Spacer,
)
from reportlab.lib.enums import TA_LEFT
from reportlab.lib import colors
import os

from db.connect import get_business_units

# Get the current directory path
current_dir = os.path.dirname(os.path.abspath(__file__))

# Path to the fonts folder
fonts_dir = os.path.join(current_dir, "..", "fonts")

DOCS_FOLDER = "docs"

# Path to the static folder
static_dir = os.path.join(current_dir, "..", "static")


bu_intro = {
    "digitalx": "The digital X Business Unit works with clients to innovate, deliver, and run solutions that drive growth and bring new business models across industries to the market.\n\
We manage two key products. First, msg.IoTA, which enables dynamic risk insights, insurance models for loss prevention, sustainability, and claims services. \
Second, we are the exclusive development partner for SAP Commerce Cloud, Financial Services Accelerator. \
Utilizing our products combined with market-leading partner solutions, such as SAP’s Customer Experience portfolio, we help our clients create meaningful customer journeys for their customers.\n\
Our global, multi-disciplinary scrum teams combine business process knowledge and development skills with advanced analytics and cloud technologies to solve the business digitization challenges of our clients. \
We are looking for open-minded people with a passion for technology and excellence to join our team.",
    "bts": "The Business Unit Business Technology Services is a special unit compared to other BU’s in msg global because we are domain and technology agnostic, catering to all industry and technology areas. \
The BTS Unit is looking for colleagues who want to successfully shape our clients' future in the digital age. We provide consulting, development, and application maintenance services for the EMEA and NA regions in technologies like the Java ecosystem, ReactJS, Testing Services, and Agile & DevOps for various industries. \
Our international team of experts is shaping the future of IT services.\n\
Our valued client is a leading provider of cutting-edge insurance software solutions to clients across the United States. As a well-established, mid-sized company headquartered in New York, they distinguish themselves through their commitment to innovation and customer satisfaction. \
They are dedicated to optimizing their core insurance software products to maintain a competitive edge within the industry.\n\
We are seeking highly motivated individuals to contribute to the development and enhancement of their core insurance software products. As a member of our dynamic team, you will play a critical role in addressing complex technical challenges, optimizing application performance, and delivering exceptional value to our clients.",
    "finance": "The Finance Business Unit provides consulting and implementation services that focus on digital transformation and innovative processes to improve the finance and accounting function for international insurance and banking companies. \
We are SAP’s preferred solution integrator and offer functional and technical expertise, project methodologies, accelerators and templates, system training, and ongoing maintenance support.\n\
Our Finance team covers all areas of implementation, data integration, and reporting. As one of the largest Finance teams in the world, we employ consultants, data scientists, software engineers, and industry experts.",
    "analytics": "The Analytics Business Unit provides consulting, implementation, training, architecture, and installation services for profitability and sustainability management based on SAP’s analytics solutions SAP Analytics Cloud (SAC), PaPM, Sustainability Control Tower (SCT), Data Warehouse Cloud (DWC) and others.\
We co-develop the SAP PaPM application and content packages providing standard functionality for various use cases and industries. Our team combines functional, technical, and industry expertise on cost allocations, financial planning/budgeting, intercompany transfer pricing, product/service costing, IT costing, tax calculations, funds/liquidity transfer pricing, cost-to-serve, planning simulation, and sustainability. \
We help our clients in many industries across the globe become more profitable and sustainable, increase operational transparency and control, and create a better basis for decision-making. We are a global, diverse, and inspiring team and offer the opportunity to extend personal and professional capabilities, by working for a leading strategic SAP partner.",
    "default": "TechInterrupt is a leading technology company\
specializing in software development. It is a system integrator,\
software development partner and managed services provider\
that helps companies improve their operational efficiency\
and decision-making capabilities.",
}

bu_offer = {
    "default": "A place where individuals are equally valued and where diversity and cultural differences are cherished.\n\
A global team of highly respected SAP and industry experts where you can make a difference.\n\
Competitive salaries and a broad range of benefits.",
    "bts": "A challenging and multi-cultural working environment with experienced teams.\n\
Project assignments and regular training schemes to learn and apply modern state-of-the-art technologies as well as professional systems development for critical business and enterprise solutions. \n\
Highly competitive compensation packages including incentive payment and private medical insurance.\n\
International exposure, internal and external training to help you further develop your talents.\n\
A team in which the core values are collaboration thought leadership and entrepreneurship.",
}

company_info = "msg global solutions is a systems integrator, software development partner, \
and managed services provider focused on SAP solutions for multiple industries. \
Our services include strategies for accounting, finance, regulatory reporting, \
performance management, sustainability, customer experience, and IoT. \
Operating from offices across the globe and growing, we help clients achieve operational efficiency \
and improve decision-making capabilities. \
With deep industry knowledge, technical expertise, and a diverse range of perspectives, \
our people spark change and create innovative solutions to complex operational issues.\
Our goal is to create long-lasting client relationships built on trust and dependability.\n\n\
msg global solutions is part of msg, an independent, international group of companies with more than 10,000 \
employees around the world."

footer_text = "msg global is an Equal Opportunity Employer. Equal Employment Opportunity has been, \
and will continue to be, a fundamental principle for us. At the heart of this policy is our \
commitment that we make job related decisions based on the job related criteria. More specifically, \
employment is based on personal capabilities and qualifications without discrimination based on race, \
color, religion, sex, age, national origin, disability, sexual orientation, marital status, ancestry, \
veteran status or any other protected characteristic as established by law. These principles are to be \
applied to policies and procedures relating to recruitment and hiring, compensation, benefits, termination \
and all other terms and conditions of employment."


def save_jd_and_retrieve(llm_response, bu_id):
    """Persist the JD into folder"""

    bu_name = get_business_units(bu_id)[0].get("name")
    folder_path = "".join([DOCS_FOLDER, "/", bu_name])

    pathlib.Path(folder_path).mkdir(parents=True, exist_ok=True)

    save_jd_doc(
        llm_response,
        folder_path=folder_path,
        bu_name=bu_name.lower(),
    )
    save_jd_pdf(
        llm_response,
        folder_path=folder_path,
        bu_name=bu_name.lower(),
    )
    return llm_response.job_title


def save_jd_doc(llm_response, folder_path, bu_name):
    """Persist the JD as Document"""

    file_name = llm_response.job_title

    doc = Document()

    # Get the page width from the document's section
    section = doc.sections[0]
    page_width = section.page_width
    left_margin = section.left_margin
    right_margin = section.right_margin

    # Calculate the available width for the image (page width minus margins)
    available_width = page_width - left_margin - right_margin

    # Add image and scale it to fit the page width

    doc.add_picture(
        os.path.join(static_dir, "jd-template-header.png"), width=available_width
    )

    add_styled_paragraph(
        p=doc.add_paragraph(),
        text="Job Position ID: 1",
        font_size=10,
        bold=True,
        space_before=20,
    )
    add_styled_paragraph(
        p=doc.add_paragraph(),
        text=file_name,
        font_size=16,
        bold=True,
        color=RGBColor(139, 0, 0),  # darkred
        space_before=30,
        space_after=20,
    )

    # Add label-value pairs
    add_label_value(doc, "Location:", "Bengaluru, India")
    add_label_value(doc, "Employment Type:", "Full Time")
    add_label_value(doc, "Work Model:", "Hybrid")

    add_styled_paragraph(
        p=doc.add_paragraph(),
        text="Who we are",
        font_size=16,
        color=RGBColor(139, 0, 0),  # darkred
        space_before=30,
        space_after=20,
    )
    for company_info_line in company_info.splitlines():
        add_styled_paragraph(
            p=doc.add_paragraph(),
            text=company_info_line,
        )

    add_styled_paragraph(
        p=doc.add_paragraph(),
        text="What we do",
        font_size=16,
        color=RGBColor(139, 0, 0),  # darkred
        space_before=30,
        space_after=20,
    )
    bu_introductions = (
        bu_intro.get(bu_name)
        if bu_intro.get(bu_name) is not None
        else bu_intro.get("default")
    )
    for introduction_line in bu_introductions.splitlines():
        add_styled_paragraph(p=doc.add_paragraph(), text=introduction_line)

    add_styled_paragraph(
        p=doc.add_paragraph(),
        text="What you will do",
        font_size=16,
        color=RGBColor(139, 0, 0),  # darkred
        space_before=30,
        space_after=20,
    )
    for responsibility in llm_response.responsibilities:
        add_styled_paragraph(p=doc.add_paragraph(), text=responsibility, bullet=True)

    add_styled_paragraph(
        p=doc.add_paragraph(),
        text="What we are looking for",
        font_size=16,
        color=RGBColor(139, 0, 0),  # darkred
        space_before=30,
        space_after=20,
    )
    for skill in llm_response.skills_experience:
        add_styled_paragraph(
            p=doc.add_paragraph(),
            text=skill,
            bullet=True,
        )

    add_styled_paragraph(
        p=doc.add_paragraph(),
        text="What we offer",
        font_size=16,
        color=RGBColor(139, 0, 0),  # darkred
        space_before=30,
        space_after=20,
    )
    bu_offers = (
        bu_offer.get(bu_name)
        if bu_offer.get(bu_name) is not None
        else bu_offer.get("default")
    )

    for offer in bu_offers.splitlines():
        add_styled_paragraph(
            p=doc.add_paragraph(),
            text=offer,
            bullet=True,
        )

    doc.add_picture(
        os.path.join(static_dir, "jd-template-footer.png"), width=available_width
    )

    add_styled_paragraph(
        p=doc.add_paragraph(),
        text=footer_text,
        space_before=20,
    )

    doc.save(f"{folder_path}/{file_name}.docx")
    return f"{file_name}.docx"


def save_jd_pdf(llm_response, folder_path, bu_name):
    """Persist the JD as PDF"""

    file_name = llm_response.job_title

    # Create PDF document
    doc = SimpleDocTemplate(
        f"{folder_path}/{file_name}.pdf",
        pagesize=letter,
        # rightMargin=72,
        # leftMargin=72,
        # topMargin=72,
        # bottomMargin=18,
    )

    # Register Arial and Arial-Bold fonts
    pdfmetrics.registerFont(TTFont("Arial", os.path.join(fonts_dir, "arial.ttf")))
    pdfmetrics.registerFont(
        TTFont("Arial-Bold", os.path.join(fonts_dir, "arialbd.ttf"))
    )

    # Define specific styles
    arial_style = ParagraphStyle(
        name="ArialHeading",
        fontName="Arial",
        fontSize=16,
        textColor=colors.darkred,
        leading=18,
        spaceBefore=30,
        spaceAfter=20,
    )
    arial_regular_style = ParagraphStyle(
        name="ArialRegular",
        fontName="Arial",
        fontSize=14,
        textColor=colors.black,
        leading=18,
        spaceBefore=30,
    )

    arial_regular_style_bullet = ParagraphStyle(
        name="ArialRegular",
        fontName="Arial",
        fontSize=14,
        textColor=colors.black,
        leading=18,
    )

    arial_bold_style = ParagraphStyle(
        name="ArialBold",
        fontName="Arial-Bold",
        fontSize=16,
        textColor=colors.darkred,
        spaceBefore=30,
        spaceAfter=20,
    )

    arial_bold_regular_style = ParagraphStyle(
        name="ArialBoldRegular",
        fontName="Arial-Bold",
        fontSize=10,
        textColor=colors.black,
        spaceBefore=20,
    )
    # Create a style for the combined text
    style_combined = ParagraphStyle(
        name="CombinedStyle",
        fontSize=10,
        textColor=colors.black,
        alignment=0,  # Left alignment
        spaceAfter=1,
    )

    # Content for the PDF
    content = []

    # Adding the header image
    header_img = Image(
        os.path.join(static_dir, "jd-template-header.png"), width=500, height=200
    )
    header_img.hAlign = "CENTER"
    content.append(header_img)

    content.append(Paragraph("Job Position ID: 1", style=arial_bold_regular_style))

    # Title (centered)
    title_style = arial_bold_style
    title_style.alignment = TA_LEFT
    content.append(Paragraph(file_name, title_style))
    # Define the labels and values with inline styling
    details = [
        ("Location:", "Bengaluru, India", "Arial-Bold", "Arial"),
        ("Employment Type:", "Full Time", "Arial-Bold", "Arial"),
        ("Work Model:", "Hybrid", "Arial-Bold", "Arial"),
    ]
    # Append each detail to the content
    for label, value, label_font, value_font in details:
        # Create formatted strings with inline styles
        formatted_label = f'<font name="{label_font}" size="10">{label}</font>'
        formatted_value = (
            f'<font name="{value_font}" size="10"> {value}</font>'  # Space before value
        )
        combined_text = f"{formatted_label}{formatted_value}"

        # Append the combined text as a paragraph
        content.append(Paragraph(combined_text, style_combined))

    # Company Info
    content.append(Paragraph("Who we are", style=arial_style))
    for company_info_line in company_info.splitlines():
        content.append(Paragraph(company_info_line, style=arial_regular_style))

    # BU introduction
    content.append(Paragraph("What we do", style=arial_style))
    bu_introductions = (
        bu_intro.get(bu_name)
        if bu_intro.get(bu_name) is not None
        else bu_intro.get("default")
    )
    for introduction in bu_introductions.splitlines():
        content.append(Paragraph(introduction, style=arial_regular_style))

    # Responsibilities
    content.append(Paragraph("What you will do", style=arial_style))
    responsibilities = [
        ListItem(Paragraph(resp, style=arial_regular_style_bullet))
        for resp in llm_response.responsibilities
    ]
    content.append(ListFlowable(responsibilities, bulletType="bullet"))

    # Skills & Experience
    content.append(Paragraph("What we are looking for", style=arial_style))
    skills = [
        ListItem(Paragraph(skill, style=arial_regular_style_bullet))
        for skill in llm_response.skills_experience
    ]
    content.append(ListFlowable(skills, bulletType="bullet"))

    content.append(Paragraph("What we offer", style=arial_style))
    bu_offers = (
        bu_offer.get(bu_name)
        if bu_offer.get(bu_name) is not None
        else bu_offer.get("default")
    )

    offers = [
        ListItem(Paragraph(offer, style=arial_regular_style_bullet))
        for offer in bu_offers.splitlines(".")
    ]
    content.append(ListFlowable(offers, bulletType="bullet"))

    # Adding the footer image
    footer_img = Image(
        os.path.join(static_dir, "jd-template-footer.png"), width=500, height=200
    )
    footer_img.hAlign = "CENTER"
    content.append(footer_img)

    # Add a Spacer to create space after the image
    content.append(Spacer(1, 20))  # 20 points of vertical space

    content.append(Paragraph(footer_text, style=arial_regular_style))
    # Build PDF
    try:
        doc.build(content)
        print("PDF has been created successfully.")
    except Exception as e:
        print(f"Error creating PDF: {e}")

    return f"{file_name}.pdf"


def add_styled_paragraph(
    p: Paragraph,
    text,
    font_name="Arial",
    font_size=14,
    bold=False,
    italic=False,
    color=RGBColor(0, 0, 0),
    bullet=False,
    space_before=0,
    space_after=0,
    line_spacing=1.5,
    alignment=WD_PARAGRAPH_ALIGNMENT.LEFT,
):
    # Create a new paragraph

    # Add text to the paragraph
    run = p.add_run(text)

    # Set font properties
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic

    # Set font color
    if color:
        run.font.color.rgb = RGBColor(*color)

    # Set paragraph alignment
    p.alignment = alignment

    # Set paragraph spacing
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)

    # Set line spacing
    # line spacing is in points; multiply by 12 for a line height factor
    p.paragraph_format.line_spacing = Pt(line_spacing * 12)

    # Add bullet points if specified
    if bullet:
        p.style = "ListBullet"  # Use built-in bullet style


def add_label_value(doc, label, value, label_font="Arial-Bold", value_font="Arial"):
    # Create a new paragraph
    p = doc.add_paragraph()

    # Add the label with bold style
    label_run = p.add_run(label)
    label_run.font.name = label_font
    label_run.font.size = Pt(10)
    label_run.font.bold = True  # Make the label bold
    label_run.font.color.rgb = RGBColor(0, 0, 0)

    # Add the value (after a space)
    value_run = p.add_run(f" {value}")
    value_run.font.name = value_font
    value_run.font.size = Pt(10)
    value_run.font.bold = False  # Regular text for value
    value_run.font.color.rgb = RGBColor(0, 0, 0)
