"""Save the JD"""

import pathlib
import pypandoc
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    ListFlowable,
    ListItem,
)
from reportlab.lib.enums import TA_CENTER

from db.connect import get_business_units

DOCS_FOLDER = "docs"

bu_intro = {
    "digitalx": "The digital X Business Unit works with clients to innovate, deliver, and run solutions that drive growth and bring new business models across industries to the market.\n\
We manage two key products. First, msg.IoTA, which enables dynamic risk insights, insurance models for loss prevention, sustainability, and claims services. \
Second, we are the exclusive development partner for SAP Commerce Cloud, Financial Services Accelerator. \
Utilizing our products combined with market-leading partner solutions, such as SAPs Customer Experience portfolio, we help our clients create meaningful customer journeys for their customers.\n\
Our global, multi-disciplinary scrum teams combine business process knowledge and development skills with advanced analytics and cloud technologies to solve the business digitization challenges of our clients. \
We are looking for open-minded people with a passion for technology and excellence to join our team.",
    "bts": "The Business Unit Business Technology Services is a special unit compared to other BU’s in msg global because we are domain and technology agnostic, catering to all industry and technology areas. \
The BTS Unit is looking for colleagues who want to successfully shape our clients' future in the digital age. We provide consulting, development, and application maintenance services for the EMEA and NA regions in technologies like the Java ecosystem, ReactJS, Testing Services, and Agile & DevOps for various industries. \
Our international team of experts is shaping the future of IT services.\n\
Our valued client is a leading provider of cutting-edge insurance software solutions to clients across the United States. As a well-established, mid-sized company headquartered in New York, they distinguish themselves through their commitment to innovation and customer satisfaction. \
They are dedicated to optimizing their core insurance software products to maintain a competitive edge within the industry.\n\
We are seeking highly motivated individuals to contribute to the development and enhancement of their core insurance software products. As a member of our dynamic team, you will play a critical role in addressing complex technical challenges, optimizing application performance, and delivering exceptional value to our clients.",
    "default": "TechInterrupt is a leading technology company\
specializing in software development. It is a system integrator,\
software development partner and managed services provider\
that helps companies improve their operational efficiency\
and decision-making capabilities.",
}

msg_footer = {
    "digitalx": "A place where individuals are equally valued and where diversity and cultural differences are cherished.\
A global team of highly respected SAP and industry experts where you can make a difference.\
Competitive salaries and a broad range of benefits, some of which are highlighted below.\
Add here the local specific (if any).\
Add here the local specific (if any)",
    "default": "A challenging and multi-cultural working environment with experienced teams.\
Project assignments and regular training schemes to learn and apply modern state-of-the-art technologies as well as professional systems development for critical business and enterprise solutions. \
Highly competitive compensation packages including incentive payment and private medical insurance.\
International exposure, internal and external training to help you further develop your talents.\
A team in which the core values are collaboration thought leadership and entrepreneurship.",
}


def save_jd_and_retrieve(llm_response, job_title, bu_id):
    """Persist the JD into folder"""

    bu_name = get_business_units(bu_id)[0].get("name")
    folder_path = "".join([DOCS_FOLDER, "/", bu_name])

    pathlib.Path(folder_path).mkdir(parents=True, exist_ok=True)

    save_jd_doc(
        llm_response, file_name=job_title, folder_path=folder_path, bu_name=bu_name
    )
    save_jd_pdf(
        llm_response, file_name=job_title, folder_path=folder_path, bu_name=bu_name
    )
    return job_title


def save_jd_doc(llm_response, file_name, folder_path, bu_name):
    """Persist the JD as Document"""
    bullet_style = "List Bullet"

    doc = Document()
    page_title = doc.add_heading(
        file_name,
        level=0,
    )
    page_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("What we do", level=1)
    bu_introductions = (
        bu_intro.get(bu_name)
        if bu_intro.get(bu_name) is not None
        else bu_intro.get("default")
    )
    for introduction_line in bu_introductions.splitlines():
        doc.add_paragraph(introduction_line)

    doc.add_heading("What you will do", level=1)
    for responsibility in llm_response.get("Responsibilities"):
        doc.add_paragraph(responsibility, bullet_style)

    doc.add_heading("What we are looking for", level=1)
    for responsibility in llm_response.get("SkillsAndExperience"):
        doc.add_paragraph(responsibility, bullet_style)

    doc.add_heading("What we offer", level=1)
    msg_footers = (
        msg_footer.get(bu_name)
        if msg_footer.get(bu_name) is not None
        else msg_footer.get("default")
    )
    for footer in msg_footers.split("."):
        doc.add_paragraph(footer, bullet_style)

    doc.save(f"{folder_path}/{file_name}.docx")
    return f"{file_name}.docx"


def save_jd_pdf(llm_response, file_name, folder_path, bu_name):
    """Persist the JD as PDF"""

    # Create PDF document
    doc = SimpleDocTemplate(
        f"{folder_path}/{file_name}.pdf",
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=18,
    )

    # Styles
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="Justify", alignment=4))  # 4 is for justified text
    styles.add(ParagraphStyle(name="Center", alignment=TA_CENTER))

    # Content for the PDF
    content = []

    # Title (centered)
    title_style = styles["Title"]
    title_style.alignment = TA_CENTER
    content.append(Paragraph(file_name, title_style))
    content.append(Spacer(1, 24))

    # BU introduction
    content.append(Paragraph("What we do", styles["Heading2"]))
    bu_introductions = (
        bu_intro.get(bu_name)
        if bu_intro.get(bu_name) is not None
        else bu_intro.get("default")
    )
    for introduction in bu_introductions.splitlines():
        content.append(Paragraph(introduction, styles["BodyText"]))

    content.append(Spacer(1, 12))

    # Responsibilities
    content.append(Paragraph("What you will do", styles["Heading2"]))
    responsibilities = [
        ListItem(Paragraph(resp, styles["BodyText"]))
        for resp in llm_response.get("Responsibilities")
    ]
    content.append(ListFlowable(responsibilities, bulletType="bullet"))
    content.append(Spacer(1, 12))

    # Skills & Experience
    content.append(Paragraph("What we are looking for", styles["Heading2"]))
    skills = [
        ListItem(Paragraph(skill, styles["BodyText"]))
        for skill in llm_response.get("SkillsAndExperience")
    ]
    content.append(ListFlowable(skills, bulletType="bullet"))

    # Footer
    content.append(Paragraph("What we offer", styles["Heading2"]))
    msg_footers = (
        msg_footer.get(bu_name)
        if msg_footer.get(bu_name) is not None
        else msg_footer.get("default")
    )
    footers = [
        ListItem(Paragraph(footer, styles["BodyText"]))
        for footer in msg_footers.split(".")
    ]
    content.append(ListFlowable(footers, bulletType="bullet"))

    # Build PDF
    try:
        doc.build(content)
        print("PDF has been created successfully.")
    except Exception as e:
        print(f"Error creating PDF: {e}")

    return f"{file_name}.pdf"


def save_jd_txt(llm_response, file_name, folder_path):
    """Persist the JD as TXT"""
    save_jd_doc(llm_response, file_name, folder_path)
    pypandoc.convert_file(
        f"{folder_path}/{file_name}.docx",
        "plain",
        outputfile=f"{folder_path}/{file_name}.txt",
    )
    return f"{file_name}.txt"


# Helper function to wrap text within the specified width
def wrap_text(c, text, max_width):
    lines = []
    words = text.split(" ")
    current_line = ""
    for word in words:
        if c.stringWidth(current_line + word, "Helvetica", 12) < max_width:
            current_line += word + " "
        else:
            lines.append(current_line.strip())
            current_line = word + " "
    if current_line:
        lines.append(current_line.strip())
    return lines
